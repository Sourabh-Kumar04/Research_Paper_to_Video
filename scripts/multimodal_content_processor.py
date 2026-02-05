"""
Multi-Modal Content Understanding System

This module provides comprehensive multi-modal content processing capabilities including:
- PDF parsing with figure and equation extraction
- Table-to-visualization conversion
- LaTeX equation recognition and rendering
- Automatic citation and reference linking
- Scientific content analysis and enhancement

Optimized for 16GB RAM systems with efficient memory management.
"""

import os
import re
import json
import hashlib
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass, asdict
from datetime import datetime
import asyncio
import logging

# Core libraries
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import seaborn as sns

# PDF processing
try:
    import fitz  # PyMuPDF
    import pdfplumber
    from pdf2image import convert_from_path
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing libraries not available. Install PyMuPDF, pdfplumber, and pdf2image.")

# OCR and text processing
try:
    import pytesseract
    from transformers import pipeline, AutoTokenizer, AutoModel
    import torch
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not available. Install pytesseract and transformers.")

# LaTeX processing
try:
    import sympy
    from sympy.parsing.latex import parse_latex
    import matplotlib.mathtext as mathtext
    LATEX_AVAILABLE = True
except ImportError:
    LATEX_AVAILABLE = False
    logging.warning("LaTeX processing not available. Install sympy.")

# Table processing
try:
    import camelot
    import tabula
    TABLE_EXTRACTION_AVAILABLE = True
except ImportError:
    TABLE_EXTRACTION_AVAILABLE = False
    logging.warning("Table extraction libraries not available. Install camelot-py and tabula-py.")

@dataclass
class ExtractedContent:
    """Container for extracted content from documents."""
    text: str = ""
    figures: List[Dict] = None
    tables: List[Dict] = None
    equations: List[Dict] = None
    citations: List[Dict] = None
    metadata: Dict = None
    
    def __post_init__(self):
        if self.figures is None:
            self.figures = []
        if self.tables is None:
            self.tables = []
        if self.equations is None:
            self.equations = []
        if self.citations is None:
            self.citations = []
        if self.metadata is None:
            self.metadata = {}

@dataclass
class ProcessingConfig:
    """Configuration for multi-modal content processing."""
    extract_figures: bool = True
    extract_tables: bool = True
    extract_equations: bool = True
    extract_citations: bool = True
    ocr_enabled: bool = True
    max_memory_mb: int = 4096  # 4GB max for 16GB system
    temp_dir: Optional[str] = None
    output_formats: List[str] = None
    
    def __post_init__(self):
        if self.output_formats is None:
            self.output_formats = ['png', 'svg', 'json']
        if self.temp_dir is None:
            self.temp_dir = tempfile.gettempdir()

class MultiModalContentProcessor:
    """
    Advanced multi-modal content processor for research papers and documents.
    
    Features:
    - PDF parsing with structure preservation
    - Figure and image extraction with OCR
    - Table extraction and conversion to visualizations
    - LaTeX equation recognition and rendering
    - Citation and reference linking
    - Memory-efficient processing for large documents
    """
    
    def __init__(self, config: ProcessingConfig = None):
        self.config = config or ProcessingConfig()
        self.logger = logging.getLogger(__name__)
        
        # Initialize AI models for content understanding
        self._init_ai_models()
        
        # Create temporary directory for processing
        self.temp_dir = Path(self.config.temp_dir) / "multimodal_processing"
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        
        # Statistics tracking
        self.processing_stats = {
            "documents_processed": 0,
            "figures_extracted": 0,
            "tables_extracted": 0,
            "equations_extracted": 0,
            "citations_extracted": 0
        }
    
    def _init_ai_models(self):
        """Initialize AI models for content understanding."""
        self.models = {}
        
        try:
            # Lightweight models for 16GB RAM constraint
            if OCR_AVAILABLE and torch.cuda.is_available():
                # Use GPU if available for better performance
                device = "cuda"
            else:
                device = "cpu"
            
            # Text classification for content type detection
            self.models['classifier'] = pipeline(
                "text-classification",
                model="microsoft/DialoGPT-medium",  # Lightweight model
                device=0 if device == "cuda" else -1
            )
            
            self.logger.info(f"AI models initialized on {device}")
            
        except Exception as e:
            self.logger.warning(f"Could not initialize AI models: {e}")
            self.models = {}
    
    async def process_document(self, file_path: str, output_dir: str = None) -> ExtractedContent:
        """
        Process a document and extract all multi-modal content.
        
        Args:
            file_path: Path to the document file
            output_dir: Directory to save extracted content
            
        Returns:
            ExtractedContent object with all extracted information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        if output_dir is None:
            output_dir = self.temp_dir / f"processed_{file_path.stem}"
        
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        self.logger.info(f"Processing document: {file_path}")
        
        # Initialize content container
        content = ExtractedContent()
        content.metadata = {
            "source_file": str(file_path),
            "processing_timestamp": datetime.now().isoformat(),
            "file_size": file_path.stat().st_size,
            "output_directory": str(output_dir)
        }
        
        try:
            # Determine file type and process accordingly
            if file_path.suffix.lower() == '.pdf':
                content = await self._process_pdf(file_path, output_dir, content)
            elif file_path.suffix.lower() in ['.txt', '.md']:
                content = await self._process_text(file_path, output_dir, content)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                content = await self._process_word_document(file_path, output_dir, content)
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")
            
            # Post-processing: enhance extracted content
            content = await self._enhance_extracted_content(content, output_dir)
            
            # Save processing results
            await self._save_processing_results(content, output_dir)
            
            # Update statistics
            self.processing_stats["documents_processed"] += 1
            
            self.logger.info(f"Document processing completed: {len(content.figures)} figures, "
                           f"{len(content.tables)} tables, {len(content.equations)} equations extracted")
            
            return content
            
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    async def _process_pdf(self, pdf_path: Path, output_dir: Path, content: ExtractedContent) -> ExtractedContent:
        """Process PDF document with comprehensive extraction."""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF processing libraries not available")
        
        self.logger.info(f"Processing PDF: {pdf_path}")
        
        # Open PDF with PyMuPDF for structure analysis
        pdf_doc = fitz.open(str(pdf_path))
        
        try:
            # Extract text content
            full_text = ""
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                full_text += page.get_text() + "\n"
            
            content.text = full_text
            
            # Extract figures and images
            if self.config.extract_figures:
                content.figures = await self._extract_pdf_figures(pdf_doc, output_dir)
            
            # Extract tables using multiple methods
            if self.config.extract_tables:
                content.tables = await self._extract_pdf_tables(pdf_path, output_dir)
            
            # Extract equations
            if self.config.extract_equations:
                content.equations = await self._extract_equations_from_text(full_text, output_dir)
            
            # Extract citations and references
            if self.config.extract_citations:
                content.citations = await self._extract_citations(full_text)
            
            # Update metadata
            content.metadata.update({
                "page_count": len(pdf_doc),
                "pdf_metadata": pdf_doc.metadata
            })
            
        finally:
            pdf_doc.close()
        
        return content
    
    async def _extract_pdf_figures(self, pdf_doc, output_dir: Path) -> List[Dict]:
        """Extract figures and images from PDF."""
        figures = []
        figures_dir = output_dir / "figures"
        figures_dir.mkdir(exist_ok=True)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            
            # Get images from page
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # Extract image
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        # Save image
                        img_filename = f"figure_p{page_num + 1}_{img_index + 1}.png"
                        img_path = figures_dir / img_filename
                        pix.save(str(img_path))
                        
                        # Analyze image content
                        img_analysis = await self._analyze_image_content(img_path)
                        
                        figure_info = {
                            "id": f"fig_{page_num + 1}_{img_index + 1}",
                            "filename": img_filename,
                            "path": str(img_path),
                            "page": page_num + 1,
                            "position": img_index + 1,
                            "size": {"width": pix.width, "height": pix.height},
                            "analysis": img_analysis,
                            "extracted_text": ""
                        }
                        
                        # OCR if enabled
                        if self.config.ocr_enabled and OCR_AVAILABLE:
                            try:
                                ocr_text = pytesseract.image_to_string(str(img_path))
                                figure_info["extracted_text"] = ocr_text.strip()
                            except Exception as e:
                                self.logger.warning(f"OCR failed for {img_filename}: {e}")
                        
                        figures.append(figure_info)
                        self.processing_stats["figures_extracted"] += 1
                    
                    pix = None  # Free memory
                    
                except Exception as e:
                    self.logger.warning(f"Could not extract image {img_index} from page {page_num + 1}: {e}")
        
        return figures
    
    async def _extract_pdf_tables(self, pdf_path: Path, output_dir: Path) -> List[Dict]:
        """Extract tables from PDF using multiple methods."""
        tables = []
        tables_dir = output_dir / "tables"
        tables_dir.mkdir(exist_ok=True)
        
        try:
            # Method 1: Use camelot for table extraction
            if TABLE_EXTRACTION_AVAILABLE:
                camelot_tables = camelot.read_pdf(str(pdf_path), pages='all')
                
                for i, table in enumerate(camelot_tables):
                    table_filename = f"table_{i + 1}.csv"
                    table_path = tables_dir / table_filename
                    
                    # Save as CSV
                    table.to_csv(str(table_path))
                    
                    # Create visualization
                    viz_path = await self._create_table_visualization(table.df, tables_dir / f"table_{i + 1}_viz.png")
                    
                    table_info = {
                        "id": f"table_{i + 1}",
                        "filename": table_filename,
                        "path": str(table_path),
                        "visualization_path": str(viz_path) if viz_path else None,
                        "shape": table.df.shape,
                        "confidence": table.accuracy if hasattr(table, 'accuracy') else None,
                        "data_preview": table.df.head().to_dict() if not table.df.empty else {},
                        "extraction_method": "camelot"
                    }
                    
                    tables.append(table_info)
                    self.processing_stats["tables_extracted"] += 1
            
            # Method 2: Use tabula as fallback
            try:
                tabula_tables = tabula.read_pdf(str(pdf_path), pages='all', multiple_tables=True)
                
                for i, df in enumerate(tabula_tables):
                    if df.empty:
                        continue
                    
                    table_filename = f"table_tabula_{i + 1}.csv"
                    table_path = tables_dir / table_filename
                    
                    # Save as CSV
                    df.to_csv(str(table_path), index=False)
                    
                    # Create visualization
                    viz_path = await self._create_table_visualization(df, tables_dir / f"table_tabula_{i + 1}_viz.png")
                    
                    table_info = {
                        "id": f"table_tabula_{i + 1}",
                        "filename": table_filename,
                        "path": str(table_path),
                        "visualization_path": str(viz_path) if viz_path else None,
                        "shape": df.shape,
                        "data_preview": df.head().to_dict(),
                        "extraction_method": "tabula"
                    }
                    
                    tables.append(table_info)
                    self.processing_stats["tables_extracted"] += 1
                    
            except Exception as e:
                self.logger.warning(f"Tabula table extraction failed: {e}")
        
        except Exception as e:
            self.logger.warning(f"Table extraction failed: {e}")
        
        return tables
    
    async def _create_table_visualization(self, df: pd.DataFrame, output_path: Path) -> Optional[Path]:
        """Create visualization for extracted table data."""
        try:
            # Determine best visualization type based on data
            if df.empty:
                return None
            
            plt.figure(figsize=(12, 8))
            
            # Check if data is numeric for different visualization types
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            
            if len(numeric_cols) >= 2:
                # Create correlation heatmap for numeric data
                corr_matrix = df[numeric_cols].corr()
                sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', center=0)
                plt.title('Data Correlation Matrix')
            
            elif len(numeric_cols) == 1:
                # Create histogram for single numeric column
                df[numeric_cols[0]].hist(bins=20)
                plt.title(f'Distribution of {numeric_cols[0]}')
                plt.xlabel(numeric_cols[0])
                plt.ylabel('Frequency')
            
            else:
                # Create a simple table visualization
                # Show first few rows as a table plot
                fig, ax = plt.subplots(figsize=(12, 6))
                ax.axis('tight')
                ax.axis('off')
                
                # Limit to first 10 rows and 6 columns for readability
                display_df = df.head(10).iloc[:, :6]
                table = ax.table(cellText=display_df.values,
                               colLabels=display_df.columns,
                               cellLoc='center',
                               loc='center')
                table.auto_set_font_size(False)
                table.set_fontsize(8)
                table.scale(1.2, 1.5)
                
                plt.title('Extracted Table Data')
            
            plt.tight_layout()
            plt.savefig(str(output_path), dpi=150, bbox_inches='tight')
            plt.close()
            
            return output_path
            
        except Exception as e:
            self.logger.warning(f"Could not create table visualization: {e}")
            return None
    
    async def _extract_equations_from_text(self, text: str, output_dir: Path) -> List[Dict]:
        """Extract and render LaTeX equations from text."""
        equations = []
        
        if not LATEX_AVAILABLE:
            self.logger.warning("LaTeX processing not available")
            return equations
        
        equations_dir = output_dir / "equations"
        equations_dir.mkdir(exist_ok=True)
        
        # Common LaTeX equation patterns
        equation_patterns = [
            r'\$\$(.+?)\$\$',  # Display math
            r'\$(.+?)\$',      # Inline math
            r'\\begin\{equation\}(.+?)\\end\{equation\}',  # Equation environment
            r'\\begin\{align\}(.+?)\\end\{align\}',        # Align environment
            r'\\begin\{eqnarray\}(.+?)\\end\{eqnarray\}',  # Eqnarray environment
        ]
        
        equation_id = 1
        
        for pattern in equation_patterns:
            matches = re.finditer(pattern, text, re.DOTALL | re.MULTILINE)
            
            for match in matches:
                latex_code = match.group(1).strip()
                
                if len(latex_code) < 3:  # Skip very short matches
                    continue
                
                try:
                    # Parse LaTeX equation
                    parsed_eq = parse_latex(latex_code)
                    
                    # Render equation as image
                    eq_filename = f"equation_{equation_id}.png"
                    eq_path = equations_dir / eq_filename
                    
                    # Create equation image using matplotlib
                    fig, ax = plt.subplots(figsize=(8, 2))
                    ax.text(0.5, 0.5, f'${latex_code}$', 
                           horizontalalignment='center',
                           verticalalignment='center',
                           transform=ax.transAxes,
                           fontsize=16)
                    ax.axis('off')
                    
                    plt.tight_layout()
                    plt.savefig(str(eq_path), dpi=150, bbox_inches='tight', 
                               facecolor='white', edgecolor='none')
                    plt.close()
                    
                    equation_info = {
                        "id": f"eq_{equation_id}",
                        "latex_code": latex_code,
                        "parsed_expression": str(parsed_eq),
                        "filename": eq_filename,
                        "path": str(eq_path),
                        "position_in_text": match.start(),
                        "context": text[max(0, match.start()-100):match.end()+100]
                    }
                    
                    equations.append(equation_info)
                    self.processing_stats["equations_extracted"] += 1
                    equation_id += 1
                    
                except Exception as e:
                    self.logger.warning(f"Could not process equation '{latex_code[:50]}...': {e}")
        
        return equations
    
    async def _extract_citations(self, text: str) -> List[Dict]:
        """Extract citations and references from text."""
        citations = []
        
        # Common citation patterns
        citation_patterns = [
            r'\[(\d+)\]',  # Numbered citations [1]
            r'\(([^)]+\d{4}[^)]*)\)',  # Author-year citations (Smith 2020)
            r'([A-Z][a-z]+ et al\.?,? \d{4})',  # Et al. citations
            r'([A-Z][a-z]+ and [A-Z][a-z]+,? \d{4})',  # Two author citations
        ]
        
        citation_id = 1
        
        for pattern in citation_patterns:
            matches = re.finditer(pattern, text)
            
            for match in matches:
                citation_text = match.group(1) if match.groups() else match.group(0)
                
                citation_info = {
                    "id": f"cite_{citation_id}",
                    "citation_text": citation_text,
                    "full_match": match.group(0),
                    "position": match.start(),
                    "context": text[max(0, match.start()-50):match.end()+50],
                    "pattern_type": pattern
                }
                
                citations.append(citation_info)
                self.processing_stats["citations_extracted"] += 1
                citation_id += 1
        
        return citations
    
    async def _analyze_image_content(self, image_path: Path) -> Dict:
        """Analyze image content to determine type and characteristics."""
        try:
            with Image.open(image_path) as img:
                # Basic image properties
                analysis = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "has_transparency": img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                }
                
                # Convert to RGB for analysis
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Analyze color distribution
                colors = img.getcolors(maxcolors=256*256*256)
                if colors:
                    analysis["dominant_colors"] = len(colors)
                    analysis["is_grayscale"] = len(set(r for r, g, b in [color[1] for color in colors[:10]])) == 1
                
                # Estimate content type based on characteristics
                width, height = img.size
                aspect_ratio = width / height
                
                if aspect_ratio > 2.0:
                    analysis["likely_type"] = "chart_or_graph"
                elif 0.8 <= aspect_ratio <= 1.2:
                    analysis["likely_type"] = "diagram_or_photo"
                else:
                    analysis["likely_type"] = "figure_or_illustration"
                
                return analysis
                
        except Exception as e:
            self.logger.warning(f"Could not analyze image {image_path}: {e}")
            return {"error": str(e)}
    
    async def _enhance_extracted_content(self, content: ExtractedContent, output_dir: Path) -> ExtractedContent:
        """Enhance extracted content with AI-powered analysis."""
        try:
            # Analyze text content for topics and themes
            if content.text and self.models.get('classifier'):
                # Split text into chunks for analysis
                text_chunks = self._split_text_into_chunks(content.text, max_length=512)
                
                topics = []
                for chunk in text_chunks[:5]:  # Analyze first 5 chunks to save memory
                    try:
                        result = self.models['classifier'](chunk)
                        topics.extend(result)
                    except Exception as e:
                        self.logger.warning(f"Text classification failed: {e}")
                        break
                
                content.metadata["detected_topics"] = topics
            
            # Link figures to relevant text sections
            content = await self._link_figures_to_text(content)
            
            # Enhance table data with statistical analysis
            content = await self._enhance_table_analysis(content)
            
            # Create content summary
            content.metadata["content_summary"] = await self._create_content_summary(content)
            
        except Exception as e:
            self.logger.warning(f"Content enhancement failed: {e}")
        
        return content
    
    def _split_text_into_chunks(self, text: str, max_length: int = 512) -> List[str]:
        """Split text into chunks for processing."""
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 > max_length:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [word]
                current_length = len(word)
            else:
                current_chunk.append(word)
                current_length += len(word) + 1
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    async def _link_figures_to_text(self, content: ExtractedContent) -> ExtractedContent:
        """Link figures to relevant text sections."""
        if not content.figures or not content.text:
            return content
        
        # Simple approach: find figure references in text
        for figure in content.figures:
            figure_refs = []
            
            # Look for common figure reference patterns
            patterns = [
                rf"Figure\s+{figure['position']}",
                rf"Fig\.?\s+{figure['position']}",
                rf"figure\s+{figure['position']}",
            ]
            
            for pattern in patterns:
                matches = re.finditer(pattern, content.text, re.IGNORECASE)
                for match in matches:
                    figure_refs.append({
                        "position": match.start(),
                        "text": match.group(0),
                        "context": content.text[max(0, match.start()-100):match.end()+100]
                    })
            
            figure["text_references"] = figure_refs
        
        return content
    
    async def _enhance_table_analysis(self, content: ExtractedContent) -> ExtractedContent:
        """Enhance table data with statistical analysis."""
        for table in content.tables:
            try:
                # Load table data
                if table["path"].endswith('.csv'):
                    df = pd.read_csv(table["path"])
                    
                    # Basic statistics
                    numeric_cols = df.select_dtypes(include=[np.number]).columns
                    if len(numeric_cols) > 0:
                        stats = df[numeric_cols].describe().to_dict()
                        table["statistics"] = stats
                    
                    # Data quality assessment
                    table["data_quality"] = {
                        "missing_values": df.isnull().sum().to_dict(),
                        "duplicate_rows": df.duplicated().sum(),
                        "data_types": df.dtypes.astype(str).to_dict()
                    }
                    
            except Exception as e:
                self.logger.warning(f"Could not enhance table analysis: {e}")
        
        return content
    
    async def _create_content_summary(self, content: ExtractedContent) -> Dict:
        """Create a comprehensive summary of extracted content."""
        summary = {
            "text_length": len(content.text),
            "word_count": len(content.text.split()) if content.text else 0,
            "figures_count": len(content.figures),
            "tables_count": len(content.tables),
            "equations_count": len(content.equations),
            "citations_count": len(content.citations),
            "processing_timestamp": datetime.now().isoformat()
        }
        
        # Add content type analysis
        if content.text:
            # Simple keyword-based content type detection
            keywords = {
                "research_paper": ["abstract", "introduction", "methodology", "results", "conclusion"],
                "technical_report": ["specification", "implementation", "requirements", "design"],
                "presentation": ["slide", "presentation", "overview", "summary"],
                "manual": ["instructions", "guide", "tutorial", "how to", "step"]
            }
            
            text_lower = content.text.lower()
            content_scores = {}
            
            for content_type, type_keywords in keywords.items():
                score = sum(1 for keyword in type_keywords if keyword in text_lower)
                content_scores[content_type] = score
            
            if content_scores:
                likely_type = max(content_scores, key=content_scores.get)
                summary["likely_content_type"] = likely_type
                summary["content_type_confidence"] = content_scores[likely_type] / len(keywords[likely_type])
        
        return summary
    
    async def _save_processing_results(self, content: ExtractedContent, output_dir: Path):
        """Save processing results to files."""
        # Save main results as JSON
        results_file = output_dir / "processing_results.json"
        
        # Convert content to serializable format
        results = asdict(content)
        
        with open(results_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False, default=str)
        
        # Save text content separately
        if content.text:
            text_file = output_dir / "extracted_text.txt"
            with open(text_file, 'w', encoding='utf-8') as f:
                f.write(content.text)
        
        # Create processing report
        report_file = output_dir / "processing_report.md"
        await self._create_processing_report(content, report_file)
    
    async def _create_processing_report(self, content: ExtractedContent, report_path: Path):
        """Create a human-readable processing report."""
        report = f"""# Multi-Modal Content Processing Report

## Document Information
- **Source**: {content.metadata.get('source_file', 'Unknown')}
- **Processing Time**: {content.metadata.get('processing_timestamp', 'Unknown')}
- **File Size**: {content.metadata.get('file_size', 0):,} bytes

## Extraction Summary
- **Text Length**: {len(content.text):,} characters
- **Word Count**: {len(content.text.split()) if content.text else 0:,} words
- **Figures Extracted**: {len(content.figures)}
- **Tables Extracted**: {len(content.tables)}
- **Equations Extracted**: {len(content.equations)}
- **Citations Found**: {len(content.citations)}

## Content Analysis
"""
        
        # Add content summary if available
        if "content_summary" in content.metadata:
            summary = content.metadata["content_summary"]
            if "likely_content_type" in summary:
                report += f"- **Likely Content Type**: {summary['likely_content_type']} "
                report += f"(confidence: {summary.get('content_type_confidence', 0):.2f})\n"
        
        # Add figures section
        if content.figures:
            report += "\n## Extracted Figures\n"
            for i, figure in enumerate(content.figures, 1):
                report += f"{i}. **{figure['filename']}** (Page {figure['page']})\n"
                if figure.get('extracted_text'):
                    report += f"   - OCR Text: {figure['extracted_text'][:100]}...\n"
                if figure.get('analysis', {}).get('likely_type'):
                    report += f"   - Type: {figure['analysis']['likely_type']}\n"
        
        # Add tables section
        if content.tables:
            report += "\n## Extracted Tables\n"
            for i, table in enumerate(content.tables, 1):
                report += f"{i}. **{table['filename']}** ({table['shape'][0]} rows × {table['shape'][1]} columns)\n"
                if table.get('statistics'):
                    report += f"   - Contains numeric data with statistical analysis\n"
        
        # Add equations section
        if content.equations:
            report += "\n## Extracted Equations\n"
            for i, equation in enumerate(content.equations, 1):
                report += f"{i}. `{equation['latex_code'][:50]}...`\n"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report)
    
    async def _process_text(self, file_path: Path, output_dir: Path, content: ExtractedContent) -> ExtractedContent:
        """Process plain text or markdown files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        content.text = text
        
        # Extract equations from text
        if self.config.extract_equations:
            content.equations = await self._extract_equations_from_text(text, output_dir)
        
        # Extract citations
        if self.config.extract_citations:
            content.citations = await self._extract_citations(text)
        
        return content
    
    async def _process_word_document(self, file_path: Path, output_dir: Path, content: ExtractedContent) -> ExtractedContent:
        """Process Word documents."""
        try:
            import docx
            
            doc = docx.Document(str(file_path))
            
            # Extract text
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            content.text = full_text
            
            # Extract images from Word document
            if self.config.extract_figures:
                figures_dir = output_dir / "figures"
                figures_dir.mkdir(exist_ok=True)
                
                for i, rel in enumerate(doc.part.rels.values()):
                    if "image" in rel.target_ref:
                        img_data = rel.target_part.blob
                        img_filename = f"word_image_{i + 1}.png"
                        img_path = figures_dir / img_filename
                        
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                        
                        figure_info = {
                            "id": f"word_img_{i + 1}",
                            "filename": img_filename,
                            "path": str(img_path),
                            "source": "word_document"
                        }
                        
                        content.figures.append(figure_info)
            
            # Extract other content types
            if self.config.extract_equations:
                content.equations = await self._extract_equations_from_text(full_text, output_dir)
            
            if self.config.extract_citations:
                content.citations = await self._extract_citations(full_text)
            
        except ImportError:
            self.logger.error("python-docx not available for Word document processing")
            raise
        except Exception as e:
            self.logger.error(f"Error processing Word document: {e}")
            raise
        
        return content
    
    def get_processing_statistics(self) -> Dict:
        """Get processing statistics."""
        return self.processing_stats.copy()
    
    def cleanup_temp_files(self):
        """Clean up temporary files."""
        try:
            import shutil
            if self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)
                self.logger.info("Temporary files cleaned up")
        except Exception as e:
            self.logger.warning(f"Could not clean up temporary files: {e}")

# Utility functions for integration with video generation system

async def process_research_paper(pdf_path: str, output_dir: str = None) -> ExtractedContent:
    """
    Convenience function to process a research paper PDF.
    
    Args:
        pdf_path: Path to the PDF file
        output_dir: Directory to save extracted content
        
    Returns:
        ExtractedContent with all extracted information
    """
    config = ProcessingConfig(
        extract_figures=True,
        extract_tables=True,
        extract_equations=True,
        extract_citations=True,
        ocr_enabled=True,
        max_memory_mb=4096  # 4GB for 16GB system
    )
    
    processor = MultiModalContentProcessor(config)
    
    try:
        content = await processor.process_document(pdf_path, output_dir)
        return content
    finally:
        processor.cleanup_temp_files()

def create_content_visualization_script(content: ExtractedContent, output_path: str) -> str:
    """
    Create a Manim script to visualize the extracted content.
    
    Args:
        content: Extracted content from document
        output_path: Path to save the Manim script
        
    Returns:
        Path to the created Manim script
    """
    script_content = f'''"""
Auto-generated Manim script for visualizing extracted content.
Generated on: {datetime.now().isoformat()}
"""

from manim import *
import numpy as np

class ContentVisualization(Scene):
    def construct(self):
        # Title
        title = Text("Document Content Overview", font_size=48)
        self.play(Write(title))
        self.wait(1)
        self.play(title.animate.to_edge(UP))
        
        # Statistics
        stats_text = [
            f"Figures: {len(content.figures)}",
            f"Tables: {len(content.tables)}",
            f"Equations: {len(content.equations)}",
            f"Citations: {len(content.citations)}"
        ]
        
        stats_group = VGroup()
        for i, stat in enumerate(stats_text):
            stat_obj = Text(stat, font_size=24)
            stat_obj.shift(DOWN * i * 0.8)
            stats_group.add(stat_obj)
        
        self.play(Write(stats_group))
        self.wait(2)
        
        # Show extracted equations if any
        if content.equations:
            self.play(FadeOut(stats_group))
            
            eq_title = Text("Extracted Equations", font_size=36)
            eq_title.to_edge(UP, buff=1)
            self.play(Write(eq_title))
            
            for i, equation in enumerate(content.equations[:3]):  # Show first 3 equations
                try:
                    eq_tex = MathTex(equation['latex_code'])
                    eq_tex.shift(DOWN * (i - 1) * 1.5)
                    self.play(Write(eq_tex))
                    self.wait(1)
                except:
                    # Fallback for invalid LaTeX
                    eq_text = Text(f"Equation {i+1}: {equation['latex_code'][:30]}...", font_size=20)
                    eq_text.shift(DOWN * (i - 1) * 1.5)
                    self.play(Write(eq_text))
                    self.wait(1)
        
        self.wait(2)
'''
    
    with open(output_path, 'w') as f:
        f.write(script_content)
    
    return output_path

if __name__ == "__main__":
    # Example usage
    import asyncio
    
    async def main():
        # Example processing
        config = ProcessingConfig()
        processor = MultiModalContentProcessor(config)
        
        # Process a sample document (replace with actual path)
        # content = await processor.process_document("sample_paper.pdf")
        # print(f"Processed document with {len(content.figures)} figures and {len(content.tables)} tables")
        
        print("Multi-modal content processor initialized successfully")
        print("Available features:")
        print(f"- PDF processing: {PDF_AVAILABLE}")
        print(f"- OCR capabilities: {OCR_AVAILABLE}")
        print(f"- LaTeX processing: {LATEX_AVAILABLE}")
        print(f"- Table extraction: {TABLE_EXTRACTION_AVAILABLE}")
    
    asyncio.run(main())